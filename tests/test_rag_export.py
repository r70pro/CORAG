"""
Unit tests for the OLMOCR RAG export subsystem.

Tests:
- Case name extraction and cleaning
- History-to-pairs parsing
- Markdown export generation and contents
- Plain text export generation and markdown stripping
- Timeline CSV export parsing and extraction
"""

import os
import unittest
import csv

# Prevent system operations during import
os.environ["TESTING"] = "true"

from rag_export import (
    EXPORT_DIR,
    _extract_case_name,
    _history_to_pairs,
    _make_export_filename,
    export_chat_markdown,
    export_chat_text,
    export_timeline_csv,
    export_chat_docx,
    export_timeline_docx,
)


class TestRAGExport(unittest.TestCase):

    def setUp(self):
        # We ensure the export directory exists
        os.makedirs(EXPORT_DIR, exist_ok=True)
        self.created_files = []
        
        import sys
        from unittest.mock import MagicMock
        self.docx_mocked = False
        try:
            import docx  # noqa: F401
        except ImportError:
            self.docx_mocked = True
            self.saved_modules = {}
            for m in ["docx", "docx.shared", "docx.enum.text", "docx.enum.table", "docx.oxml", "docx.oxml.ns"]:
                if m in sys.modules:
                    self.saved_modules[m] = sys.modules[m]
                sys.modules[m] = MagicMock()
            
            # Setup Document mock
            self.mock_doc = MagicMock()
            def save_mock(path):
                with open(path, "wb") as f:
                    f.write(b"mock docx content")
            self.mock_doc.save.side_effect = save_mock
            sys.modules["docx"].Document.return_value = self.mock_doc

    def tearDown(self):
        # Clean up any files created during the test
        for path in self.created_files:
            if os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
        
        # Restore sys.modules if mocked
        if getattr(self, "docx_mocked", False):
            import sys
            for m in ["docx", "docx.shared", "docx.enum.text", "docx.enum.table", "docx.oxml", "docx.oxml.ns"]:
                if m in self.saved_modules:
                    sys.modules[m] = self.saved_modules[m]
                elif m in sys.modules:
                    del sys.modules[m]

    def track_file(self, path):
        if path:
            self.created_files.append(path)
        return path

    def test_extract_case_name(self):
        self.assertEqual(_extract_case_name("All Cases"), "all_cases")
        self.assertEqual(_extract_case_name(""), "all_cases")
        self.assertEqual(_extract_case_name(None), "all_cases")
        self.assertEqual(_extract_case_name("workspace/run_case123"), "case123")
        self.assertEqual(_extract_case_name("case@name#special!"), "case_name_special_")
        # Long name truncation
        long_name = "a" * 100
        truncated = _extract_case_name(long_name)
        self.assertEqual(len(truncated), 60)
        self.assertEqual(truncated, "a" * 60)

    def test_make_export_filename(self):
        filename = _make_export_filename("test", "md")
        self.assertTrue(filename.startswith("test_"))
        self.assertTrue(filename.endswith(".md"))

    def test_history_to_pairs_empty(self):
        self.assertEqual(_history_to_pairs(None), [])
        self.assertEqual(_history_to_pairs([]), [])

    def test_history_to_pairs_dicts(self):
        history = [
            {"role": "user", "content": "What is the diagnosis?"},
            {"role": "assistant", "content": "The diagnosis is a SLAP tear."},
        ]
        pairs = _history_to_pairs(history)
        self.assertEqual(len(pairs), 2)
        self.assertEqual(pairs[0], ("user", "What is the diagnosis?"))
        self.assertEqual(pairs[1], ("assistant", "The diagnosis is a SLAP tear."))

    def test_history_to_pairs_tuples_lists(self):
        history = [
            ("Hello", None),
            (None, "Hi there"),
            ["Question", "Answer"],
        ]
        pairs = _history_to_pairs(history)
        self.assertEqual(len(pairs), 3)
        self.assertEqual(pairs[0], ("user", "Hello"))
        self.assertEqual(pairs[1], ("assistant", "Hi there"))
        self.assertEqual(pairs[2], ("user", "Question"))

    def test_history_to_pairs_malformed(self):
        history = [
            "not a dict or list",
            {"role": "user"},  # missing content
            [123],             # too short list
        ]
        pairs = _history_to_pairs(history)
        # It should ignore malformed entries or handle them safely
        # {"role": "user"} returns ("user", "")
        # The string and [123] are skipped
        self.assertEqual(len(pairs), 1)
        self.assertEqual(pairs[0], ("user", ""))

    def test_export_chat_markdown_empty(self):
        self.assertIsNone(export_chat_markdown(None))
        self.assertIsNone(export_chat_markdown([]))

    def test_export_chat_markdown_valid(self):
        history = [
            {"role": "user", "content": "Help me."},
            {"role": "assistant", "content": "I am helping."},
        ]
        path = self.track_file(export_chat_markdown(history, mode="timeline", active_case="case1"))
        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        self.assertIn("# RAG Analysis Export", content)
        self.assertIn("**Mode:** timeline", content)
        self.assertIn("## 👤 User Query", content)
        self.assertIn("Help me.", content)
        self.assertIn("## 🤖 Analysis Response", content)
        self.assertIn("I am helping.", content)

    def test_export_chat_text_empty(self):
        self.assertIsNone(export_chat_text(None))
        self.assertIsNone(export_chat_text([]))

    def test_export_chat_text_valid(self):
        history = [
            {"role": "user", "content": "Help me."},
            {
                "role": "assistant",
                "content": "# Heading 1\nThis is **bold** and *italic* text. Check [Google](https://google.com)."
            },
        ]
        path = self.track_file(export_chat_text(history, mode="timeline", active_case="case2"))
        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        self.assertIn("RAG ANALYSIS EXPORT", content)
        self.assertIn("USER QUERY:", content)
        self.assertIn("ANALYSIS RESPONSE:", content)
        # Check formatting is stripped
        self.assertIn("Heading 1", content)
        self.assertNotIn("# Heading 1", content)
        self.assertIn("This is bold and italic text.", content)
        self.assertNotIn("**bold**", content)
        self.assertIn("Check Google.", content)
        self.assertNotIn("[Google](https://google.com)", content)

    def test_export_timeline_csv_empty(self):
        self.assertIsNone(export_timeline_csv(None))
        self.assertIsNone(export_timeline_csv([]))

    def test_export_timeline_csv_no_tables(self):
        history = [
            {"role": "user", "content": "Help me."},
            {"role": "assistant", "content": "No table here."},
        ]
        path = self.track_file(export_timeline_csv(history))
        self.assertIsNone(path)

    def test_export_timeline_csv_valid_table(self):
        table_markdown = """
Here is the timeline:

| Date | Event | Provider |
|---|---|---|
| 2020-01-01 | Injury | Employer |
| 2020-01-02 | Surgery | Dr Smith |

End of timeline.
"""
        history = [
            {"role": "user", "content": "Give me a table"},
            {"role": "assistant", "content": table_markdown},
        ]
        path = self.track_file(export_timeline_csv(history, active_case="case3"))
        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

        # Read CSV file and verify content
        rows = []
        with open(path, "r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)

        self.assertEqual(len(rows), 3)  # Header + 2 data rows
        self.assertEqual(rows[0], ["Date", "Event", "Provider"])
        self.assertEqual(rows[1], ["2020-01-01", "Injury", "Employer"])
        self.assertEqual(rows[2], ["2020-01-02", "Surgery", "Dr Smith"])

    def test_export_timeline_csv_preserves_escaped_pipes_in_provenance(self):
        history = [
            {
                "role": "assistant",
                "content": (
                    "| Date | Event | Source |\n"
                    "|---|---|---|\n"
                    "| 2020-01-01 | Injury | record.pdf \\| Page: 2 \\| Author: Dr Smith |"
                ),
            }
        ]

        path = self.track_file(export_timeline_csv(history, active_case="case-pipes"))
        with open(path, "r", encoding="utf-8", newline="") as exported:
            rows = list(csv.reader(exported))

        self.assertEqual({len(row) for row in rows}, {3})
        self.assertEqual(rows[1][2], "record.pdf | Page: 2 | Author: Dr Smith")

    def test_export_chat_docx_empty(self):
        self.assertIsNone(export_chat_docx(None))
        self.assertIsNone(export_chat_docx([]))

    def test_export_chat_docx_valid(self):
        history = [
            {"role": "user", "content": "Help me."},
            {
                "role": "assistant",
                "content": "# Heading 1\n## Heading 2\n### Heading 3\n---\n| Date | Inconsistency |\n|---|---|\n| 2020-01-01 | Patient history inconsistency |\n\n- Bullet 1\n* Bullet 2\n\nDegenerative **changes** *noted*."
            },
        ]
        path = self.track_file(export_chat_docx(history, mode="timeline", active_case="case4"))
        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

    def test_export_timeline_docx_empty(self):
        self.assertIsNone(export_timeline_docx(None))
        self.assertIsNone(export_timeline_docx([]))

    def test_export_timeline_docx_no_tables(self):
        history = [
            {"role": "user", "content": "Help me."},
            {"role": "assistant", "content": "No table here."},
        ]
        path = self.track_file(export_timeline_docx(history))
        self.assertIsNone(path)

    def test_export_timeline_docx_valid(self):
        table_markdown = """
Here is the timeline:

| Date | Event | Provider |
|---|---|---|
| 2020-01-01 | Injury | Employer |
| 2020-01-02 | Surgery | Dr Smith |

End of timeline.
"""
        history = [
            {"role": "user", "content": "Give me a table"},
            {"role": "assistant", "content": table_markdown},
        ]
        path = self.track_file(export_timeline_docx(history, active_case="case5"))
        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

    def test_export_timeline_docx_no_headers_no_rows(self):
        # Empty/false table format test to cover lines 494->485
        history = [
            {"role": "user", "content": "Empty table"},
            {"role": "assistant", "content": "|\n|-|"},
        ]
        path = self.track_file(export_timeline_docx(history, active_case="case6"))
        self.assertIsNone(path)

    def test_docx_import_errors(self):
        # Simulate ImportError inside docx helpers when docx components are missing (sys.modules None)
        import sys
        from unittest.mock import patch
        
        # We patch sys.modules inside a block to force ImportError
        with patch.dict(sys.modules, {
            "docx.shared": None,
            "docx.enum.text": None,
            "docx.oxml": None,
            "docx.oxml.ns": None,
            "docx.enum.table": None,
            "docx": None
        }):
            from rag_export import _add_letterhead, _set_bottom_border, _add_docx_table
            # Call helpers to ensure they fail gracefully/return None
            # Since docx is None, Document() raises or import fails. But let's check functions
            self.assertIsNone(_add_letterhead(None))
            self.assertIsNone(_set_bottom_border(None))
            self.assertIsNone(_add_docx_table(None, ["H"], [["R"]]))
            self.assertIsNone(export_chat_docx([{"role": "user", "content": "H"}]))
            self.assertIsNone(export_timeline_docx([{"role": "user", "content": "H"}]))

    def test_firm_logo_handling(self):
        import docx
        import rag_export
        from unittest.mock import patch
        
        doc = docx.Document()
        
        # Test add_picture success
        with patch("os.path.exists", return_value=True):
            with patch.object(doc, "add_picture") as mock_add:
                with patch("rag_export.FIRM_LOGO", "dummy_logo.png"):
                    rag_export._add_letterhead(doc)
                    mock_add.assert_called_once_with("dummy_logo.png", width=unittest.mock.ANY)
                    
        # Test add_picture exception (lines 258-262)
        with patch("os.path.exists", return_value=True):
            with patch.object(doc, "add_picture", side_effect=Exception("failed to add logo")):
                with patch("rag_export.FIRM_LOGO", "dummy_logo.png"):
                    # This should run without raising exception
                    rag_export._add_letterhead(doc)

    def test_add_docx_table_empty(self):
        import docx
        from rag_export import _add_docx_table
        doc = docx.Document()
        # Direct call to empty headers and rows (line 380)
        self.assertIsNone(_add_docx_table(doc, [], []))
