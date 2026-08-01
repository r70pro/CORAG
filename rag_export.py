"""
RAG Export — Generate downloadable files from chat analysis sessions.

Provides:
- Markdown export of full chat sessions
- Plain text export for paste into Word/email
- CSV export of timeline-mode table output
- DOCX export with firm letterhead (court-ready analysis reports)
"""

import csv
import datetime
import os
import re

from path_security import resolve_file_under, resolve_under
from settings_manager import WORKSPACE_DIR

EXPORT_DIR = str(resolve_under(WORKSPACE_DIR, "exports"))

# ── Letterhead configuration ─────────────────────────────────
# Court-ready exports are branded with a firm letterhead. Override these via
# environment variables (OLMOCR_FIRM_NAME, OLMOCR_FIRM_SUBTITLE,
# OLMOCR_FIRM_CONTACT) or edit defaults here. A logo PNG at the path in
# OLMOCR_FIRM_LOGO (if set) is embedded at the top of the document.
FIRM_NAME = os.environ.get("OLMOCR_FIRM_NAME", "Your Firm & Associates")
FIRM_SUBTITLE = os.environ.get("OLMOCR_FIRM_SUBTITLE", "Medicolegal Document Analysis")
FIRM_CONTACT = os.environ.get(
    "OLMOCR_FIRM_CONTACT", "Level 1, 123 Example Street · Telephone (00) 0000 0000"
)
FIRM_LOGO = os.environ.get("OLMOCR_FIRM_LOGO", "")


def _ensure_export_dir():
    """Create the exports directory if needed."""
    resolve_under(WORKSPACE_DIR, "exports").mkdir(parents=True, exist_ok=True)


def _make_export_filename(prefix, ext):
    """Generate a timestamped export filename."""
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_{ts}.{ext}"


def _extract_case_name(active_case):
    """Extract a human-readable case name from the active case label."""
    if not active_case or active_case == "All Cases":
        return "all_cases"
    # Strip common prefixes and clean for filenames
    name = active_case.replace("workspace/", "").replace("run_", "")
    name = re.sub(r"[^\w\-]", "_", name)
    return name[:60]


def _history_to_pairs(history, include_reasoning=False):
    """Convert Gradio chatbot history (list of dicts) to (role, content) pairs."""
    pairs = []
    if not history:
        return pairs
    for msg in history:
        if isinstance(msg, dict):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
        elif isinstance(msg, list | tuple) and len(msg) >= 2:
            role = "user" if msg[0] is not None else "assistant"
            content = msg[0] if msg[0] else msg[1]
        else:
            continue
        pairs.append((role, str(content)))
        if include_reasoning and isinstance(msg, dict) and msg.get("reasoning"):
            pairs.append(("reasoning", str(msg["reasoning"])))
    return pairs


def _parse_markdown_table_row(line):
    """Parse one pipe-delimited Markdown row, preserving escaped pipes.

    LLM-generated provenance commonly contains ``\\|`` inside a cell. A plain
    ``str.split("|")`` corrupts those citations and produces invalid CSV/DOCX
    column counts.
    """
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return []
    if not stripped[1:-1].strip():
        return []

    cells = []
    current = []
    escaped = False
    for char in stripped[1:-1]:
        if escaped:
            if char == "|":
                current.append("|")
            else:
                current.extend(("\\", char))
            escaped = False
        elif char == "\\":
            escaped = True
        elif char == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def export_chat_markdown(history, mode="Free Q&A", active_case="All Cases", include_reasoning=False):
    """Export a chat session as a Markdown file.

    Args:
        history: Gradio chatbot history (list of message dicts).
        mode: Current analysis mode name.
        active_case: Active case label.

    Returns:
        File path to the generated .md file, or None on error.
    """
    if not history:
        return None

    _ensure_export_dir()
    case_name = _extract_case_name(active_case)
    filename = _make_export_filename(f"analysis_{case_name}", "md")
    filepath = resolve_file_under(EXPORT_DIR, filename, {".md"})

    pairs = _history_to_pairs(history, include_reasoning=include_reasoning)

    lines = [
        "# RAG Analysis Export",
        "",
        f"**Mode:** {mode}",
        f"**Case:** {active_case}",
        f"**Exported:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**Messages:** {len(pairs)}",
        "",
        "---",
        "",
    ]

    for role, content in pairs:
        if role == "user":
            lines.append("## 👤 User Query")
            lines.append("")
            lines.append(content)
            lines.append("")
        elif role == "reasoning":
            lines.append("## 🛡️ Administrative LLM Reasoning Audit")
            lines.append("")
            lines.append(content)
            lines.append("")
            lines.append("---")
            lines.append("")
        else:
            lines.append("## 🤖 Analysis Response")
            lines.append("")
            lines.append(content)
            lines.append("")
            lines.append("---")
            lines.append("")

    with filepath.open("w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return str(filepath)


def export_chat_text(history, mode="Free Q&A", active_case="All Cases", include_reasoning=False):
    """Export a chat session as a plain text file.

    Args:
        history: Gradio chatbot history (list of message dicts).
        mode: Current analysis mode name.
        active_case: Active case label.

    Returns:
        File path to the generated .txt file, or None on error.
    """
    if not history:
        return None

    _ensure_export_dir()
    case_name = _extract_case_name(active_case)
    filename = _make_export_filename(f"analysis_{case_name}", "txt")
    filepath = resolve_file_under(EXPORT_DIR, filename, {".txt"})

    pairs = _history_to_pairs(history, include_reasoning=include_reasoning)

    lines = [
        "RAG ANALYSIS EXPORT",
        f"{'=' * 60}",
        f"Mode:      {mode}",
        f"Case:      {active_case}",
        f"Exported:  {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Messages:  {len(pairs)}",
        f"{'=' * 60}",
        "",
    ]

    for role, content in pairs:
        if role == "user":
            lines.append("USER QUERY:")
            lines.append(f"{'-' * 40}")
            lines.append(content)
            lines.append("")
        elif role == "reasoning":
            lines.append("ADMINISTRATIVE LLM REASONING AUDIT:")
            lines.append(f"{'-' * 40}")
            lines.append(content)
            lines.append("")
        else:
            lines.append("ANALYSIS RESPONSE:")
            lines.append(f"{'-' * 40}")
            # Strip markdown formatting for plain text
            plain = content
            # Remove markdown headers
            plain = re.sub(r"^#{1,6}\s+", "", plain, flags=re.MULTILINE)
            # Remove bold/italic markers
            plain = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", plain)
            # Remove markdown links but keep text
            plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", plain)
            lines.append(plain)
            lines.append("")
            lines.append(f"{'=' * 60}")
            lines.append("")

    with filepath.open("w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return str(filepath)


def export_timeline_csv(history, active_case="All Cases"):
    """Extract timeline table from chat history and export as CSV.

    Parses Markdown tables from assistant responses in timeline mode.

    Args:
        history: Gradio chatbot history (list of message dicts).
        active_case: Active case label.

    Returns:
        File path to the generated .csv file, or None on error/no tables.
    """
    if not history:
        return None

    _ensure_export_dir()
    case_name = _extract_case_name(active_case)
    filename = _make_export_filename(f"timeline_{case_name}", "csv")
    filepath = resolve_file_under(EXPORT_DIR, filename, {".csv"})

    # Extract all markdown tables from assistant messages
    table_rows = []
    headers = None

    pairs = _history_to_pairs(history)

    for role, content in pairs:
        if role != "assistant" or not content:
            continue

        # Find markdown table lines: lines starting with |
        lines = content.split("\n")
        in_table = False
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("|") and stripped.endswith("|"):
                # Skip separator lines (e.g., |---|---|---|)
                if re.match(r"^\|[\s\-:|]+\|$", stripped):
                    continue
                # Parse cells
                cells = _parse_markdown_table_row(stripped)
                if not headers:
                    headers = cells
                    in_table = True
                else:
                    table_rows.append(cells)
                    in_table = True
            elif in_table and not stripped.startswith("|"):
                in_table = False

    if not headers and not table_rows:
        return None

    with filepath.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        for row in table_rows:
            writer.writerow(row)

    return str(filepath)


# ── DOCX export (firm letterhead) ────────────────────────────


def _add_letterhead(doc):
    """Add a branded firm letterhead to the top of a DOCX document."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        return

    if FIRM_LOGO and os.path.exists(FIRM_LOGO):
        try:
            doc.add_picture(FIRM_LOGO, width=Pt(120))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(FIRM_NAME)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(FIRM_SUBTITLE)
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(FIRM_CONTACT)
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Horizontal rule via bottom border on a blank paragraph
    rule = doc.add_paragraph()
    _set_bottom_border(rule)


def _set_bottom_border(paragraph):
    """Apply a bottom border to a paragraph (used as a horizontal rule)."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F2937")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _render_markdown_to_docx(doc, text):
    """Render a Markdown string into the DOCX document as paragraphs/tables.

    Supports headings (## / ###), fenced tables, horizontal rules, bold/italic,
    and bullet lists. Not a full Markdown parser — sufficient for analysis
    output. Falls back to plain paragraphs for anything unrecognised.
    """
    from docx.shared import Pt

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Markdown table (starts with | and next line is a separator)
        if (
            line.lstrip().startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1])
        ):
            headers = [c.strip() for c in line.strip().split("|")[1:-1]]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().split("|")[1:-1]])
                i += 1
            _add_docx_table(doc, headers, rows)
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(13)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(15)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(17)
        elif line.startswith("---"):
            _set_bottom_border(doc.add_paragraph())
        elif re.match(r"^\s*[-*]\s+", line):
            item = re.sub(r"^\s*[-*]\s+", "", line)
            doc.add_paragraph(_inline_md(item), style="List Bullet")
        else:
            doc.add_paragraph(_inline_md(line))
        i += 1


def _inline_md(text):
    """Convert inline **bold** and *italic* markers to runs; returns plain text
    with markers stripped if python-docx styling is not desired inline."""
    # python-docx paragraphs accept plain strings; we strip markdown markers so
    # the exported document is clean rather than showing raw asterisks.
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    return text


def _add_docx_table(doc, headers, rows):
    """Add a styled table to the DOCX document."""
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return
    if not headers and not rows:
        return
    table = doc.add_table(rows=1, cols=max(len(headers), 1))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr[idx].text = _inline_md(h)
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            if idx < len(cells):
                cells[idx].text = _inline_md(val)


def export_chat_docx(history, mode="Free Q&A", active_case="All Cases", include_reasoning=False):
    """Export a full chat session as a branded DOCX report.

    The report carries the firm letterhead and renders each user query and
    analysis response as styled content.

    Args:
        history: Gradio chatbot history (list of message dicts).
        mode: Current analysis mode name.
        active_case: Active case label.

    Returns:
        File path to the generated .docx file, or None on error.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        return None

    if not history:
        return None

    _ensure_export_dir()
    case_name = _extract_case_name(active_case)
    filename = _make_export_filename(f"analysis_{case_name}", "docx")
    filepath = resolve_file_under(EXPORT_DIR, filename, {".docx"})

    doc = Document()
    _add_letterhead(doc)

    title = doc.add_paragraph()
    tr = title.add_run("Medicolegal Analysis Report")
    tr.bold = True
    tr.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Mode: {mode}    |    Case: {active_case}    |    "
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ).italic = True

    pairs = _history_to_pairs(history, include_reasoning=include_reasoning)
    for role, content in pairs:
        if role == "user":
            label = doc.add_paragraph()
            lr = label.add_run("User Query")
            lr.bold = True
            lr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            doc.add_paragraph(_inline_md(content))
        elif role == "reasoning":
            label = doc.add_paragraph()
            lr = label.add_run("Administrative LLM Reasoning Audit")
            lr.bold = True
            lr.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
            _render_markdown_to_docx(doc, content)
        else:
            label = doc.add_paragraph()
            lr = label.add_run("Analysis Response")
            lr.bold = True
            lr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            _render_markdown_to_docx(doc, content)
        _set_bottom_border(doc.add_paragraph())

    doc.save(str(filepath))
    return str(filepath)


def export_timeline_docx(history, active_case="All Cases"):
    """Export timeline tables from chat history as a branded DOCX document.

    Args:
        history: Gradio chatbot history (list of message dicts).
        active_case: Active case label.

    Returns:
        File path to the generated .docx file, or None on error/no tables.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return None

    if not history:
        return None

    pairs = _history_to_pairs(history)
    tables = []  # list of (headers, rows)
    for role, content in pairs:
        if role != "assistant" or not content:
            continue
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            if (
                lines[i].strip().startswith("|")
                and i + 1 < len(lines)
                and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1].strip())
            ):
                headers = _parse_markdown_table_row(lines[i])
                i += 2
                rows = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    rows.append(_parse_markdown_table_row(lines[i]))
                    i += 1
                if headers or rows:
                    tables.append((headers, rows))
            else:
                i += 1

    if not tables:
        return None

    _ensure_export_dir()
    case_name = _extract_case_name(active_case)
    filename = _make_export_filename(f"timeline_{case_name}", "docx")
    filepath = resolve_file_under(EXPORT_DIR, filename, {".docx"})

    doc = Document()
    _add_letterhead(doc)
    title = doc.add_paragraph()
    tr = title.add_run("Clinical Timeline")
    tr.bold = True
    tr.font.size = Pt(16)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Case: {active_case}    |    Generated: "
        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ).italic = True

    for headers, rows in tables:
        _add_docx_table(doc, headers, rows)
        doc.add_paragraph()

    doc.save(str(filepath))
    return str(filepath)
